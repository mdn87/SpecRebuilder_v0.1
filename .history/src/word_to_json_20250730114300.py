#!/usr/bin/env python3
"""
Word Document to JSON Converter

This script converts Word documents (.docx) to JSON format to analyze
the document structure and multilist level formatting.
"""

import json
import sys
import os
from pathlib import Path
from typing import Dict, List, Any, Optional
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET
import zipfile

class WordToJsonConverter:
    """Converts Word documents to JSON format for analysis"""
    
    def __init__(self):
        self.document_data = {}
    
    def extract_paragraph_info(self, paragraph) -> Dict[str, Any]:
        """Extract detailed information from a paragraph"""
        info = {
            'text': paragraph.text.strip(),
            'style_name': paragraph.style.name if paragraph.style else None,
            'alignment': str(paragraph.alignment) if paragraph.alignment else None,
            'runs': []
        }
        
        # Extract run information
        for run in paragraph.runs:
            run_info = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name if run.font.name else None,
                'font_size': run.font.size.pt if run.font.size else None,
                'font_color': run.font.color.rgb if run.font.color.rgb else None
            }
            info['runs'].append(run_info)
        
        # Extract numbering information
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            num_pr = paragraph._p.pPr.numPr
            info['numbering'] = {
                'id': num_pr.numId.val if num_pr.numId is not None else None,
                'level': num_pr.ilvl.val if num_pr.ilvl is not None else None
            }
        
        return info
    
    def extract_document_structure(self, docx_path: str) -> Dict[str, Any]:
        """Extract the complete document structure"""
        doc = Document(docx_path)
        
        document_info = {
            'file_path': docx_path,
            'paragraphs': [],
            'sections': [],
            'headers': [],
            'footers': [],
            'comments': [],
            'metadata': {
                'core_properties': {},
                'app_properties': {}
            }
        }
        
        # Extract paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # Only include non-empty paragraphs
                para_info = self.extract_paragraph_info(paragraph)
                para_info['index'] = i
                document_info['paragraphs'].append(para_info)
        
        # Extract sections
        for section in doc.sections:
            section_info = {
                'start_type': str(section.start_type),
                'page_width': section.page_width.inches if section.page_width else None,
                'page_height': section.page_height.inches if section.page_height else None,
                'left_margin': section.left_margin.inches if section.left_margin else None,
                'right_margin': section.right_margin.inches if section.right_margin else None,
                'top_margin': section.top_margin.inches if section.top_margin else None,
                'bottom_margin': section.bottom_margin.inches if section.bottom_margin else None,
                'header_distance': section.header_distance.inches if section.header_distance else None,
                'footer_distance': section.footer_distance.inches if section.footer_distance else None,
            }
            document_info['sections'].append(section_info)
        
        # Extract headers and footers
        for i, section in enumerate(doc.sections):
            if section.header:
                header_info = {
                    'section_index': i,
                    'paragraphs': []
                }
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        header_info['paragraphs'].append(self.extract_paragraph_info(paragraph))
                document_info['headers'].append(header_info)
            
            if section.footer:
                footer_info = {
                    'section_index': i,
                    'paragraphs': []
                }
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        footer_info['paragraphs'].append(self.extract_paragraph_info(paragraph))
                document_info['footers'].append(footer_info)
        
        # Extract comments
        try:
            comments = self.extract_comments(docx_path)
            document_info['comments'] = comments
        except Exception as e:
            document_info['comments'] = []
            print(f"Warning: Could not extract comments: {e}")
        
        # Extract metadata
        try:
            document_info['metadata']['core_properties'] = {
                'title': doc.core_properties.title,
                'subject': doc.core_properties.subject,
                'creator': doc.core_properties.creator,
                'created': str(doc.core_properties.created) if doc.core_properties.created else None,
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else None,
                'last_modified_by': doc.core_properties.last_modified_by,
                'revision': doc.core_properties.revision,
                'keywords': doc.core_properties.keywords,
                'category': doc.core_properties.category,
                'comments': doc.core_properties.comments,
                'language': doc.core_properties.language,
            }
        except Exception as e:
            print(f"Warning: Could not extract core properties: {e}")
        
        return document_info
    
    def extract_comments(self, docx_path: str) -> List[Dict[str, Any]]:
        """Extract comments from the document"""
        comments = []
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                if 'word/comments.xml' in zip_file.namelist():
                    comments_xml = zip_file.read('word/comments.xml')
                    root = ET.fromstring(comments_xml)
                    
                    for comment in root.findall('.//w:comment', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        comment_info = {
                            'id': comment.get('id'),
                            'author': comment.get('author'),
                            'date': comment.get('date'),
                            'text': ''
                        }
                        
                        # Extract comment text
                        for text_elem in comment.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            comment_info['text'] += text_elem.text or ''
                        
                        comments.append(comment_info)
        except Exception as e:
            print(f"Warning: Could not extract comments: {e}")
        
        return comments
    
    def convert_to_json(self, docx_path: str, output_path: Optional[str] = None) -> str:
        """Convert Word document to JSON"""
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Document not found: {docx_path}")
        
        # Extract document structure
        document_data = self.extract_document_structure(docx_path)
        
        # Determine output path
        if output_path is None:
            base_name = Path(docx_path).stem
            output_path = f"{base_name}_structure.json"
        
        # Save to JSON
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(document_data, f, indent=2, ensure_ascii=False, default=str)
        
        print(f"Document converted to JSON: {output_path}")
        print(f"Extracted {len(document_data['paragraphs'])} paragraphs")
        print(f"Found {len(document_data['sections'])} sections")
        print(f"Found {len(document_data['headers'])} headers")
        print(f"Found {len(document_data['footers'])} footers")
        print(f"Found {len(document_data['comments'])} comments")
        
        return output_path

def main():
    """Main function"""
    if len(sys.argv) < 2:
        print("Usage: python word_to_json.py <docx_file> [output_file]")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    converter = WordToJsonConverter()
    try:
        output_file = converter.convert_to_json(docx_path, output_path)
        print(f"Successfully converted {docx_path} to {output_file}")
    except Exception as e:
        print(f"Error converting document: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main() 