#!/usr/bin/env python3
"""
Preserving Word Document Sanitizer

This script preserves list structure while fixing XML issues that cause
"unreadable content" warnings. It extracts the content and numbering,
then rebuilds it with proper Word structure.
"""

import os
import sys
import json
import zipfile
import tempfile
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def extract_numbering_info(docx_path: str):
    """Extract numbering information from the original document"""
    numbering_info = []
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as zipf:
            # Extract document.xml
            if 'word/document.xml' in zipf.namelist():
                doc_xml = zipf.read('word/document.xml').decode('utf-8')
                
                # Parse paragraphs and their numbering
                import xml.etree.ElementTree as ET
                root = ET.fromstring(doc_xml)
                
                for p in root.findall('.//w:p'):
                    para_info = {'text': '', 'has_numbering': False, 'level': None, 'num_id': None}
                    
                    # Get text content
                    text_elements = p.findall('.//w:t')
                    para_info['text'] = ' '.join([t.text or '' for t in text_elements])
                    
                    # Check for numbering
                    num_pr = p.find('.//w:numPr')
                    if num_pr is not None:
                        para_info['has_numbering'] = True
                        
                        # Get numbering ID
                        num_id = num_pr.find('.//w:numId')
                        if num_id is not None:
                            para_info['num_id'] = num_id.get('w:val')
                        
                        # Get level
                        ilvl = num_pr.find('.//w:ilvl')
                        if ilvl is not None:
                            para_info['level'] = int(ilvl.get('w:val'))
                    
                    if para_info['text'].strip():
                        numbering_info.append(para_info)
    
    except Exception as e:
        print(f"Error extracting numbering info: {e}")
    
    return numbering_info

def create_numbered_paragraph(doc, text, level=0):
    """Create a paragraph with proper numbering"""
    p = doc.add_paragraph()
    
    # Add numbering properties
    p._p.get_or_add_pPr().append(parse_xml(f'''
        <w:numPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:numId w:val="1"/>
            <w:ilvl w:val="{level}"/>
        </w:numPr>
    '''))
    
    # Add text
    p.add_run(text)
    
    return p

def create_numbering_definitions(doc):
    """Create numbering definitions in the document"""
    # This creates a basic multilevel list definition
    # You might need to customize this based on your specific numbering needs
    
    # Add numbering definitions to the document
    numbering_part = doc.part.get_or_add_part('word/numbering.xml')
    
    numbering_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:abstractNum w:abstractNumId="0">
        <w:lvl w:ilvl="0">
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="720" w:hanging="360"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="1440" w:hanging="360"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="2">
            <w:numFmt w:val="lowerLetter"/>
            <w:lvlText w:val="%3."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="2160" w:hanging="360"/>
            </w:pPr>
        </w:lvl>
    </w:abstractNum>
    <w:num w:numId="1">
        <w:abstractNumId w:val="0"/>
    </w:num>
</w:numbering>'''
    
    numbering_part.blob = numbering_xml.encode('utf-8')

def sanitize_preserving_lists(input_path: str, output_path: str):
    """Sanitize document while preserving list structure"""
    try:
        print(f"Extracting numbering information from: {input_path}")
        numbering_info = extract_numbering_info(input_path)
        
        print(f"Found {len(numbering_info)} paragraphs with numbering info")
        
        # Create a new document
        doc = Document()
        
        # Create numbering definitions
        create_numbering_definitions(doc)
        
        # Add paragraphs with preserved numbering
        for para_info in numbering_info:
            if para_info['has_numbering']:
                level = para_info.get('level', 0)
                create_numbered_paragraph(doc, para_info['text'], level)
            else:
                # Add regular paragraph
                doc.add_paragraph(para_info['text'])
        
        print(f"Saving sanitized document: {output_path}")
        doc.save(output_path)
        
        print("Document sanitization with preserved lists complete!")
        return True
        
    except Exception as e:
        print(f"Error sanitizing document: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Main function"""
    if len(sys.argv) < 3:
        print("Usage: python preserving_sanitizer.py <input_docx> <output_docx>")
        print("Example: python preserving_sanitizer.py output/word_compatible_output.docx output/word_compatible_output_preserved.docx")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    if not os.path.exists(input_path):
        print(f"Error: Input document not found: {input_path}")
        sys.exit(1)
    
    # Create output directory if it doesn't exist
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    
    success = sanitize_preserving_lists(input_path, output_path)
    if not success:
        sys.exit(1)

if __name__ == "__main__":
    main() 