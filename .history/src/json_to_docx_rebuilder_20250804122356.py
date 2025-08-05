#!/usr/bin/env python3
"""
JSON to DOCX Rebuilder

This script takes the JSON analysis data and creates a properly formatted
Word document using python-docx, preserving list structure and avoiding
"unreadable content" warnings.
"""

import os
import sys
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def load_json_analysis(json_path: str):
    """Load the JSON analysis data"""
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def create_numbered_paragraph(doc, text, level=0):
    """Create a paragraph with proper numbering"""
    p = doc.add_paragraph()
    
    # Add numbering properties
    p._p.get_or_add_pPr().append(parse_xml(f'''
        <w:numPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:numId w:val="1"/>
            <w:ilvl w:val="{level}"/>
        </w:numPr>
    '''))
    
    # Add text
    p.add_run(text)
    
    return p

def rebuild_document_from_json(json_path: str, output_path: str):
    """Rebuild document from JSON analysis using python-docx"""
    try:
        print(f"Loading JSON analysis from: {json_path}")
        json_data = load_json_analysis(json_path)
        
        # Create a new document
        doc = Document()
        
        # Get paragraphs from JSON
        paragraphs = json_data.get('all_paragraphs', [])
        print(f"Found {len(paragraphs)} paragraphs to process")
        
        # Add paragraphs with preserved numbering
        for para_data in paragraphs:
            text = para_data.get('text', '').strip()
            if not text:
                continue
            
            # Check if this paragraph has numbering
            has_numbering = bool(para_data.get('list_number') or para_data.get('inferred_number'))
            
            if has_numbering:
                # Use cleaned content if available, otherwise use original text
                content = para_data.get('cleaned_content', text)
                level = para_data.get('level', 0)
                create_numbered_paragraph(doc, content, level)
            else:
                # Add regular paragraph
                doc.add_paragraph(text)
        
        print(f"Saving document: {output_path}")
        doc.save(output_path)
        
        print("Document rebuild complete!")
        return True
        
    except Exception as e:
        print(f"Error rebuilding document: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Main function"""
    if len(sys.argv) < 3:
        print("Usage: python json_to_docx_rebuilder.py <json_file> <output_docx>")
        print("Example: python json_to_docx_rebuilder.py output/SECTION_00_00_00_hybrid_analysis.json output/rebuilt_from_json.docx")
        sys.exit(1)
    
    json_path = sys.argv[1]
    output_path = sys.argv[2]
    
    if not os.path.exists(json_path):
        print(f"Error: JSON file not found: {json_path}")
        sys.exit(1)
    
    # Create output directory if it doesn't exist
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    
    success = rebuild_document_from_json(json_path, output_path)
    if not success:
        sys.exit(1)

if __name__ == "__main__":
    main() 