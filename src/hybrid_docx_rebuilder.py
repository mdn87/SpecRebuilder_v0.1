#!/usr/bin/env python3
"""
Hybrid DOCX Rebuilder

This script combines python-docx for proper Word structure with
custom numbering definitions to preserve complex list hierarchies.
"""

import os
import sys
import json
import zipfile
import tempfile
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement, qn

def load_json_analysis(json_path: str):
    """Load the JSON analysis data"""
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def create_custom_numbering_xml():
    """Create custom numbering XML that matches our original structure"""
    numbering_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:abstractNum w:abstractNumId="0">
        <w:lvl w:ilvl="0">
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="720" w:hanging="360"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="1440" w:hanging="360"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="2">
            <w:numFmt w:val="lowerLetter"/>
            <w:lvlText w:val="%3."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="2160" w:hanging="360"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="3">
            <w:numFmt w:val="lowerLetter"/>
            <w:lvlText w:val="%4."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="2880" w:hanging="360"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="4">
            <w:numFmt w:val="lowerLetter"/>
            <w:lvlText w:val="%5."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="3600" w:hanging="360"/>
            </w:pPr>
        </w:lvl>
    </w:abstractNum>
    <w:num w:numId="1">
        <w:abstractNumId w:val="0"/>
    </w:num>
</w:numbering>'''
    return numbering_xml

def create_numbered_paragraph(doc, text, level=0):
    """Create a paragraph with proper numbering"""
    p = doc.add_paragraph()
    
    # Add numbering properties
    p._p.get_or_add_pPr().append(parse_xml(f'''
        <w:numPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:numId w:val="1"/>
            <w:ilvl w:val="{level}"/>
        </w:numPr>
    '''))
    
    # Add text
    p.add_run(text)
    
    return p

def inject_numbering_xml(docx_path: str, numbering_xml: str):
    """Inject custom numbering XML into the document"""
    with zipfile.ZipFile(docx_path, 'a') as zipf:
        # Check if numbering.xml already exists
        if 'word/numbering.xml' in zipf.namelist():
            # Remove existing numbering.xml
            zipf.writestr('word/numbering.xml', numbering_xml)
        else:
            # Add numbering.xml
            zipf.writestr('word/numbering.xml', numbering_xml)
        
        # Update [Content_Types].xml to include numbering
        if '[Content_Types].xml' in zipf.namelist():
            content_types = zipf.read('[Content_Types].xml').decode('utf-8')
            if 'word/numbering.xml' not in content_types:
                # Insert numbering override before the closing </Types>
                content_types = content_types.replace(
                    '</Types>',
                    '  <Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>\n</Types>'
                )
                zipf.writestr('[Content_Types].xml', content_types)

def rebuild_document_from_json(json_path: str, output_path: str):
    """Rebuild document from JSON analysis using hybrid approach"""
    try:
        print(f"Loading JSON analysis from: {json_path}")
        json_data = load_json_analysis(json_path)
        
        # Create a new document
        doc = Document()
        
        # Get paragraphs from JSON
        paragraphs = json_data.get('all_paragraphs', [])
        print(f"Found {len(paragraphs)} paragraphs to process")
        
        # Add paragraphs with preserved numbering
        for para_data in paragraphs:
            text = para_data.get('text', '').strip()
            if not text:
                continue
            
            # Check if this paragraph has numbering
            has_numbering = bool(para_data.get('list_number') or para_data.get('inferred_number'))
            
            if has_numbering:
                # Use cleaned content if available, otherwise use original text
                content = para_data.get('cleaned_content', text)
                level = para_data.get('level', 0)
                create_numbered_paragraph(doc, content, level)
            else:
                # Add regular paragraph
                doc.add_paragraph(text)
        
        print(f"Saving document: {output_path}")
        doc.save(output_path)
        
        # Inject custom numbering XML
        print("Injecting custom numbering definitions...")
        numbering_xml = create_custom_numbering_xml()
        inject_numbering_xml(output_path, numbering_xml)
        
        print("Document rebuild complete!")
        return True
        
    except Exception as e:
        print(f"Error rebuilding document: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Main function"""
    if len(sys.argv) < 3:
        print("Usage: python hybrid_docx_rebuilder.py <json_file> <output_docx>")
        print("Example: python hybrid_docx_rebuilder.py output/SECTION_00_00_00_hybrid_analysis.json output/hybrid_rebuilt.docx")
        sys.exit(1)
    
    json_path = sys.argv[1]
    output_path = sys.argv[2]
    
    if not os.path.exists(json_path):
        print(f"Error: JSON file not found: {json_path}")
        sys.exit(1)
    
    # Create output directory if it doesn't exist
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    
    success = rebuild_document_from_json(json_path, output_path)
    if not success:
        sys.exit(1)

if __name__ == "__main__":
    main() 